import os
import numpy as np
import pandas as pd
import glob
import matplotlib.pyplot as plt
from os import chdir
import sklearn
from sklearn import metrics, datasets, neighbors
import numpy as np
import matplotlib.pyplot as plt
from sklearn.datasets import make_classification
from sklearn.ensemble import RandomForestClassifier
from sklearn.inspection import permutation_importance
import sys
import warnings
from sklearn.model_selection import train_test_split
import itertools
from sklearn.metrics import plot_confusion_matrix
from sklearn.metrics import accuracy_score
from sklearn.metrics import balanced_accuracy_score
from sklearn.metrics import cohen_kappa_score
from sklearn.metrics import classification_report
from sklearn.metrics import roc_curve
from sklearn.metrics import average_precision_score
from sklearn.metrics import RocCurveDisplay
from sklearn.metrics import precision_recall_curve
from sklearn.metrics import PrecisionRecallDisplay
import numpy as np
import matplotlib.pyplot as plt

from sklearn import svm, datasets
from sklearn.model_selection import train_test_split
from sklearn.metrics import plot_confusion_matrix
from sklearn.model_selection import learning_curve
from sklearn.svm import SVC
from sklearn import svm
from ML import *
from ML import plot_confusion
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import io


warnings.filterwarnings("ignore")
np.set_printoptions(precision=2)

# Global variables are always in caps, to distinguish them from local variables.
WORKING_DIRECTORY = os.path.dirname(__file__)
DATA_DIR = os.path.join(WORKING_DIRECTORY, 'Data')
INFO_DIR = os.path.join(WORKING_DIRECTORY, 'Info')
PLOT_DIR = os.path.join(INFO_DIR, "PLOT")
Data = os.path.join(DATA_DIR, "ML_Data.csv")
ValData = os.path.join(DATA_DIR, "ML_Data_val.csv")

# Flags that change the behaviour of the control_script

# flag to demonstrate the principle of flags
Import = True
ImportVal = True
DataInvestigation = False
split = True
featureSelection = False
printSelected = False
printNotSelected = False
scaler = True
Alg_tests = True
plotAlg = True
crossVal = True
confusionMatrix = True
RobustML = False
def plot_confusion(classifier, X_test, y_test, class_names):
    # https://scikit-learn.org/stable/auto_examples/model_selection/plot_confusion_matrix.html#sphx-glr-auto-examples-model-selection-plot-confusion-matrix-py

    np.set_printoptions(precision=2)

    # Plot non-normalized confusion matrix
    titles_options = [("Confusion matrix, without normalization", None),
                      ("Normalized confusion matrix", 'true')]
    for title, normalize in titles_options:
        disp = plot_confusion_matrix(classifier, X_test, y_test,
                                     display_labels=class_names,
                                     cmap=plt.cm.Blues,
                                     normalize=normalize)
        disp.ax_.set_title(title)

        print(title)
        print(disp.confusion_matrix)

    plt.show()

def get_X_y(df):
    X = df.loc[:, df.columns != 'Conditions']
    X = X.loc[:, X.columns != "Unnamed: 0"]
    y = df.loc[:, df.columns == 'Conditions']
    return X, y

def get_age_files(x_f, filename, remove_duplicates=True):
    x = pd.read_csv(x_f)
    y_n = filename.replace("X", "y")
    y = pd.read_csv(y_n)

    if remove_duplicates:
        x=x.drop(columns="Unnamed: 0").drop_duplicates()
        y = x.join(y)
        y=y["Conditions"]
        y= pd.DataFrame(y)
    else:
        y=y.drop(columns="Unnamed: 0")
        x = x.drop(columns="Unnamed: 0")

    return train_test_split(x,y,test_size=0.2)
    

def open_files(f, filename):
    X_train = pd.read_csv(f)
    X_train = X_train.drop(columns="Unnamed: 0")

    y_train_v = filename.replace("X", "y")
    y_train = pd.read_csv(y_train_v)
    y_train = y_train.drop(columns="Unnamed: 0")

    X_test_v = filename.replace("train", "validate")
    X_test = pd.read_csv(X_test_v)
    X_test = X_test.drop(columns="Unnamed: 0")

    y_test_v = filename.replace("X_train", "y_validate")
    y_test = pd.read_csv(y_test_v)
    y_test = y_test.drop(columns="Unnamed: 0")
    return X_train, X_test, y_train, y_test

def permutation_based_feature_importance(clf, X, y, feature_names, save=False, filename = False):
    result = permutation_importance(clf, X, y, n_repeats=10, random_state=0,
                                    n_jobs=-1)

    fig, ax = plt.subplots()
    sorted_idx = result.importances_mean.argsort()
    ax.boxplot(result.importances[sorted_idx].T,
               vert=False, labels=feature_names[sorted_idx])
    ax.set_title("Permutation Importance of each feature")
    ax.set_ylabel("Features")
    fig.tight_layout()
    if not save:
        plt.show()
    else:
        if not filename:
            raise FileNotFoundError ("Provide a filename to save the file")
        else:
            loc = r"../../Early Detection/Data/Figures/"
            plt.savefig((loc+filename))

def evaluate_model(y_pred, y_true, X_test, y_test, clf, target_names, document=None):
    ######################################################
    # accuracy

    print("Accuracy: ", accuracy_score(y_true, y_pred))

    ###################################################
    # balanced accuracy
    #
    print("Balanced accuracy score: ", balanced_accuracy_score(y_true, y_pred))

    #########################
    # cohen_kappa_score
    """
    The kappa score is a number between -1 and 1. Scores above .8 are generally considered good agreement; zero or lower means no agreement (practically random labels)
    """
    print("cohen kappa score: ",cohen_kappa_score(y_true, y_pred), "above 0.8 is good agreement")

    ##############################
    # plot confusion matrix
    plot_confusion(clf, X_test, y_test, ["HD", "WT"])
    ####################################
    # classification report

    print("classification report: \n", classification_report(y_true, y_pred, target_names=target_names))
    #########################################
    # idk

    print("Precision: ",metrics.precision_score(y_true, y_pred, average="binary", pos_label="HD"))

    print("Recall:", metrics.recall_score(y_true, y_pred, average="binary", pos_label="HD"))

    print("F1:",metrics.f1_score(y_true, y_pred, average="binary", pos_label="HD"))

    print("F beta, beta-0.5", metrics.fbeta_score(y_true, y_pred, beta=0.5,average="binary", pos_label="HD"))

    print("F beta, beta-1",metrics.fbeta_score(y_true, y_pred, beta=1,average="binary", pos_label="HD"))

    print("F beta, beta-2",metrics.fbeta_score(y_true, y_pred, beta=2,average="binary", pos_label="HD"))

    print("precision recall fscore support", metrics.precision_recall_fscore_support(y_true, y_pred, beta=0.5,average="binary", pos_label="HD"))


    # ROC curve
    y_scores = clf.predict_proba(X_test)[:, 1]
    precision, recall, threshold = precision_recall_curve(y_true, y_scores, pos_label="HD")


    print("Average precision score: ", average_precision_score(y_true, y_scores, pos_label="HD"))

    #######################################
    # ROC
    # https://scikit-learn.org/stable/auto_examples/miscellaneous/plot_display_object_visualization.html#sphx-glr-auto-examples-miscellaneous-plot-display-object-visualization-py
    y_score = clf.decision_function(X_test)

    fpr, tpr, _ = roc_curve(y_test, y_score, pos_label=clf.classes_[1])
    roc_display = RocCurveDisplay(fpr=fpr, tpr=tpr).plot()

    ##################################
    # precision recall curve

    prec, recall, _ = precision_recall_curve(y_test, y_score,
                                             pos_label=clf.classes_[1])
    pr_display = PrecisionRecallDisplay(precision=prec,     recall=recall).plot()

    # combine plots
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 8))

    roc_display.plot(ax=ax1)
    pr_display.plot(ax=ax2)
    plt.show()

    if document is not None:
        document.add_heading("Test Metrics", level=2)
        document.add_paragraph(("Accuracy: {}".format(accuracy_score(y_true, y_pred))), style = "List Bullet")
        document.add_paragraph(("Balanced accuracy score: {}".format(balanced_accuracy_score(y_true, y_pred))), style = "List Bullet")
        document.add_paragraph(("Cohen kappa score: {} ".format(accuracy_score(y_true, y_pred))), style = "List Bullet")
        p=document.add_paragraph("(The kappa score is a number between -1 and 1. Scores above .8 are generally considered good agreement; zero or lower means no agreement (practically random labels))", style = "List Bullet")
        p.add_run('italic.').italic = True


        # confusion matricies
        document.add_heading("Confusion Matrices", level=2)
        memfile = io.BytesIO()
        np.set_printoptions(precision=2)

        # Plot non-normalized confusion matrix
        titles_options = [("Confusion matrix, without normalization", None),
                          ("Normalized confusion matrix", 'true')]
        for title, normalize in titles_options:
            disp = plot_confusion_matrix(clf, X_test, y_test,
                                         display_labels=["HD", "WT"],
                                         cmap=plt.cm.Blues,
                                         normalize=normalize)
            disp.ax_.set_title(title)

            plt.savefig(memfile)
            document.add_picture(memfile, width=Inches(5))
        memfile.close()

        # classification report
        document.add_heading("Classification report", level=2)
        document.add_paragraph("{}".format(classification_report(y_true, y_pred, target_names=target_names)))

        # Precision/recall
        document.add_heading("Precision/Recall Scores", level=2)
        document.add_paragraph("Precision: {}".format(metrics.precision_score(y_true, y_pred, average="binary", pos_label="HD")), style= "List Bullet")
        document.add_paragraph("Recall: {}".format(metrics.recall_score(y_true, y_pred, average="binary", pos_label="HD")), style= "List Bullet")
        document.add_paragraph("F1 {}".format(metrics.f1_score(y_true, y_pred, average="binary", pos_label="HD")), style= "List Bullet")
        document.add_paragraph("Precision Recall F-Score Support: {}".format(metrics.precision_recall_fscore_support(y_true, y_pred, beta=0.5,average="binary", pos_label="HD")))


        memfile = io.BytesIO()
        y_score = clf.decision_function(X_test)

        # precision recall curve
        prec, recall, _ = precision_recall_curve(y_test, y_score,
                                                 pos_label=clf.classes_[1])
        pr_display = PrecisionRecallDisplay(precision=prec,     recall=recall).plot()

        # combine plots
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 8))

        roc_display.plot(ax=ax1)
        pr_display.plot(ax=ax2)
        plt.savefig(memfile)
        document.add_picture(memfile, width=Inches(5))
        memfile.close()

        document.save(r'../../ML/Classifiers/demo.docx')