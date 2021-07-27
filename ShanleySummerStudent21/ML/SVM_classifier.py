
from ML import *
from docx import Document


loc = r"../InputForML/SMOTE/"
chdir(loc)

def classifySVM(f, filename, save_document=True):
    """
    classifies svm learner
    :param f:
    :param filename:
    :param save_document:
    :return: returns more vars if save_doc is true, take care when calling function
    """
    clf = svm.SVC(gamma=0.001, C=100., probability=True)

    X_train, X_test, y_train, y_test = get_age_files(f, filename, remove_duplicates=False)

    clf.fit(X_train, y_train)
    y_preds = clf.predict(X_test)

    rna = filename.replace("X_", "").replace(".csv", "")

    if save_document:
        document = Document()
        document.add_heading("Classifier Performance on {} data".format(filename.replace("X_", "").replace("_", " ").replace("train.csv", ""), "classifier"),0)
        document.add_heading("Training", level=2)
        document.add_paragraph(('Training Accuracy : %.3f'%clf.score(X_train, y_train)))
        return clf, X_test, y_test, y_preds, ["HD", "WT"], document, rna, X_train, y_train

    print(filename,'\nTraining Accuracy : %.3f'%clf.score(X_train, y_train))
    return clf, X_test, y_test, y_preds, ["HD", "WT"], rna,X_train, y_train



""":cvar



for filename in glob.glob('*X*'):
    with open(os.path.join(os.getcwd(), filename), 'r') as f:
        clf, X_test, y_test, y_preds, classes , document, name,X_train, y_train = classifySVM(f, filename)
        evaluate_model(y_preds, y_test, X_test, y_test, clf, classes, X_train, y_train, document=document, fname=name.replace("_train", ""))

"""